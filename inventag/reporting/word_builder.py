#!/usr/bin/env python3
"""
Word Document Builder

Word document builder for professional BOM generation.
Enhanced with professional formatting and template system.

Features:
- Word document generation with professional formatting and template system
- Executive summary section with compliance overview and key metrics
- Service-specific resource tables with custom descriptions and formatting
- Network analysis section with CIDR utilization details and diagrams
- Security analysis section with risk assessment summaries and recommendations
- Table of contents generation and cross-reference management
"""

import logging
import os
from typing import Dict, List, Any, Optional, Set
from datetime import datetime, timezone
from collections import defaultdict

from .document_generator import DocumentBuilder, DocumentGenerationResult, BOMData

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    # Create placeholders for type hints when python-docx is not available
    Document = Any
    RGBColor = Any
    Inches = Any
    Pt = Any
    WD_ALIGN_PARAGRAPH = Any
    WD_TABLE_ALIGNMENT = Any
    WD_ALIGN_VERTICAL = Any


class WordDocumentBuilder(DocumentBuilder):
    """
    Word document builder for professional BOM generation.
    
    Enhanced with professional formatting and template system:
    - Word document generation with professional formatting and template system
    - Executive summary section with compliance overview and key metrics
    - Service-specific resource tables with custom descriptions and formatting
    - Network analysis section with CIDR utilization details and diagrams
    - Security analysis section with risk assessment summaries and recommendations
    - Table of contents generation and cross-reference management
    """
    
    def __init__(self, config):
        super().__init__(config)
        
        # Word styling configuration
        self.styles = self._initialize_styles()
        
    def can_handle_format(self, format_type: str) -> bool:
        """Check if this builder can handle Word format."""
        return format_type.lower() == "word"
        
    def validate_dependencies(self) -> List[str]:
        """Validate Word dependencies."""
        if not PYTHON_DOCX_AVAILABLE:
            return ["python-docx library not available - install with: pip install python-docx"]
        return []
        
    def _initialize_styles(self) -> Dict[str, Any]:
        """Initialize Word styling configuration."""
        if not PYTHON_DOCX_AVAILABLE:
            return {}
            
        return {
            "title_font_size": Pt(18),
            "heading1_font_size": Pt(16),
            "heading2_font_size": Pt(14),
            "heading3_font_size": Pt(12),
            "body_font_size": Pt(11),
            "table_font_size": Pt(10),
            "primary_color": self._hex_to_rgb(self.config.branding.color_scheme.get("primary", "366092")),
            "accent_color": self._hex_to_rgb(self.config.branding.color_scheme.get("accent", "70AD47")),
            "danger_color": self._hex_to_rgb(self.config.branding.color_scheme.get("danger", "C5504B")),
            "warning_color": self._hex_to_rgb(self.config.branding.color_scheme.get("warning", "FFC000")),
            "font_name": self.config.branding.font_family
        }
        
    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor object."""
        if not PYTHON_DOCX_AVAILABLE:
            return None
            
        # Remove # if present
        hex_color = hex_color.lstrip('#')
        
        # Handle 3-character hex codes
        if len(hex_color) == 3:
            hex_color = ''.join([c*2 for c in hex_color])
            
        # Default to black if invalid
        if len(hex_color) != 6:
            hex_color = "000000"
            
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except ValueError:
            return RGBColor(0, 0, 0)  # Default to black
            
    def generate_document(self, bom_data: BOMData, output_path: str) -> DocumentGenerationResult:
        """
        Generate Word document with professional formatting.
        
        Enhanced with professional formatting and template system:
        - Word document generation with professional formatting and template system
        - Executive summary section with compliance overview and key metrics
        - Service-specific resource tables with custom descriptions and formatting
        - Network analysis section with CIDR utilization details and diagrams
        - Security analysis section with risk assessment summaries and recommendations
        - Table of contents generation and cross-reference management
        """
        start_time = datetime.now(timezone.utc)
        
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return DocumentGenerationResult(
                    format_type="word",
                    filename=os.path.basename(output_path) if output_path else "",
                    success=False,
                    error_message="python-docx library not available"
                )
                
            self.logger.info(f"Generating Word document: {output_path}")
            
            # Create document
            doc = Document()
            
            # Configure document styles
            self._configure_document_styles(doc)
            
            # Add document content
            self._add_title_page(doc, bom_data)
            self._add_table_of_contents(doc)
            self._add_executive_summary(doc, bom_data)
            self._add_service_sections(doc, bom_data)
            self._add_network_analysis_section(doc, bom_data)
            self._add_security_analysis_section(doc, bom_data)
            self._add_compliance_details_section(doc, bom_data)
            self._add_appendices(doc, bom_data)
            
            # Apply branding
            self._apply_document_branding(doc)
            
            # Save document
            doc.save(output_path)
            
            end_time = datetime.now(timezone.utc)
            
            self.logger.info(f"Word document generated successfully: {output_path}")
            
            return DocumentGenerationResult(
                format_type="word",
                filename=os.path.basename(output_path),
                success=True,
                generation_time_seconds=(end_time - start_time).total_seconds()
            )
            
        except Exception as e:
            self.logger.error(f"Word generation failed: {e}")
            return DocumentGenerationResult(
                format_type="word",
                filename=os.path.basename(output_path) if output_path else "",
                success=False,
                error_message=str(e)
            )
            
    def _configure_document_styles(self, doc: Document):
        """Configure document styles for professional formatting."""
        if not PYTHON_DOCX_AVAILABLE:
            return
            
        styles = doc.styles
        
        # Configure Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.styles["font_name"]
        normal_font.size = self.styles["body_font_size"]
        
        # Configure heading styles
        try:
            heading1_style = styles['Heading 1']
            heading1_font = heading1_style.font
            heading1_font.name = self.styles["font_name"]
            heading1_font.size = self.styles["heading1_font_size"]
            heading1_font.color.rgb = self.styles["primary_color"]
            heading1_font.bold = True
            
            heading2_style = styles['Heading 2']
            heading2_font = heading2_style.font
            heading2_font.name = self.styles["font_name"]
            heading2_font.size = self.styles["heading2_font_size"]
            heading2_font.color.rgb = self.styles["primary_color"]
            heading2_font.bold = True
            
            heading3_style = styles['Heading 3']
            heading3_font = heading3_style.font
            heading3_font.name = self.styles["font_name"]
            heading3_font.size = self.styles["heading3_font_size"]
            heading3_font.bold = True
            
        except KeyError:
            # Styles might not exist in all templates
            self.logger.warning("Some heading styles not available in document template")
            
    def _add_title_page(self, doc: Document, bom_data: BOMData):
        """Add professional title page."""
        # Title
        title = doc.add_heading(f"{self.config.branding.company_name}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = self.styles["title_font_size"]
        title_run.font.color.rgb = self.styles["primary_color"]
        
        # Subtitle
        subtitle = doc.add_heading("Cloud Bill of Materials Report", level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = self.styles["primary_color"]
        
        # Add some spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata table
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Table Grid'
        metadata_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table width
        metadata_table.autofit = False
        for column in metadata_table.columns:
            column.width = Inches(2.5)
            
        # Add metadata
        metadata = bom_data.generation_metadata
        
        cells = metadata_table.rows[0].cells
        cells[0].text = "Generated Date:"
        cells[1].text = metadata.get("generated_at", "Unknown")
        
        cells = metadata_table.rows[1].cells
        cells[0].text = "Total Resources:"
        cells[1].text = str(metadata.get("total_resources", 0))
        
        cells = metadata_table.rows[2].cells
        cells[0].text = "Processing Time:"
        cells[1].text = f"{metadata.get('processing_time_seconds', 0):.2f} seconds"
        
        cells = metadata_table.rows[3].cells
        cells[0].text = "Report Version:"
        cells[1].text = "1.0"
        
        # Format metadata table
        for row in metadata_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = self.styles["body_font_size"]
                        
        # Make first column bold
        for row in metadata_table.rows:
            first_cell = row.cells[0]
            for paragraph in first_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Page break
        doc.add_page_break()
        
    def _add_table_of_contents(self, doc: Document):
        """Add table of contents placeholder."""
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Add TOC entries manually (python-docx doesn't support automatic TOC)
        toc_entries = [
            "Executive Summary",
            "Service Resources",
            "Network Analysis",
            "Security Analysis", 
            "Compliance Details",
            "Appendices"
        ]
        
        for entry in toc_entries:
            toc_para = doc.add_paragraph()
            toc_run = toc_para.add_run(f"• {entry}")
            toc_run.font.size = self.styles["body_font_size"]
            
        doc.add_page_break()
        
    def _add_executive_summary(self, doc: Document, bom_data: BOMData):
        """
        Add executive summary section with compliance overview and key metrics.
        
        Enhanced executive summary with:
        - Compliance overview and key metrics
        - Service breakdown summary
        - Key findings and recommendations
        """
        # Section heading
        summary_heading = doc.add_heading("Executive Summary", level=1)
        summary_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Overview paragraph
        overview_para = doc.add_paragraph()
        overview_text = (
            f"This Cloud Bill of Materials (BOM) report provides a comprehensive analysis of "
            f"AWS resources across the {self.config.branding.company_name} infrastructure. "
            f"The report covers {bom_data.generation_metadata.get('total_resources', 0)} resources "
            f"and includes compliance analysis, network security assessment, and operational recommendations."
        )
        overview_run = overview_para.add_run(overview_text)
        overview_run.font.size = self.styles["body_font_size"]
        
        # Compliance overview
        doc.add_heading("Compliance Overview", level=2)
        
        compliance = bom_data.compliance_summary
        total_resources = compliance.get("total_resources", 0)
        compliant_resources = compliance.get("compliant_resources", 0)
        non_compliant_resources = compliance.get("non_compliant_resources", 0)
        compliance_pct = compliance.get("compliance_percentage", 0)
        
        # Compliance summary table
        compliance_table = doc.add_table(rows=4, cols=2)
        compliance_table.style = 'Table Grid'
        
        # Add compliance data
        compliance_data = [
            ("Total Resources", str(total_resources)),
            ("Compliant Resources", str(compliant_resources)),
            ("Non-Compliant Resources", str(non_compliant_resources)),
            ("Compliance Percentage", f"{compliance_pct:.1f}%")
        ]
        
        for i, (label, value) in enumerate(compliance_data):
            row = compliance_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Format cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = self.styles["body_font_size"]
                        
            # Make first column bold
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
            # Color code compliance values
            if i == 1:  # Compliant resources
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = self.styles["accent_color"]
            elif i == 2:  # Non-compliant resources
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = self.styles["danger_color"]
                        
        # Service breakdown
        doc.add_heading("Service Breakdown", level=2)
        
        # Count resources by service
        services = defaultdict(int)
        for resource in bom_data.resources:
            service = resource.get("service", "Unknown")
            services[service] += 1
            
        # Service breakdown table
        service_table = doc.add_table(rows=len(services) + 1, cols=3)
        service_table.style = 'Table Grid'
        
        # Headers
        header_row = service_table.rows[0]
        header_row.cells[0].text = "Service"
        header_row.cells[1].text = "Resource Count"
        header_row.cells[2].text = "Percentage"
        
        # Format headers
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = self.styles["body_font_size"]
                    run.font.color.rgb = self.styles["primary_color"]
                    
        # Service data
        for i, (service, count) in enumerate(sorted(services.items()), 1):
            row = service_table.rows[i]
            percentage = (count / total_resources * 100) if total_resources > 0 else 0
            
            row.cells[0].text = service
            row.cells[1].text = str(count)
            row.cells[2].text = f"{percentage:.1f}%"
            
            # Format cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = self.styles["body_font_size"]
                        
        # Key findings
        doc.add_heading("Key Findings", level=2)
        
        findings = self._generate_key_findings(bom_data)
        for finding in findings:
            finding_para = doc.add_paragraph()
            finding_run = finding_para.add_run(f"• {finding}")
            finding_run.font.size = self.styles["body_font_size"]
            
        doc.add_page_break()
        
    def _add_service_sections(self, doc: Document, bom_data: BOMData):
        """
        Add service-specific resource tables with custom descriptions and formatting.
        
        Enhanced service sections with:
        - Service-specific resource tables with custom descriptions and formatting
        - Resource details with compliance status highlighting
        - Service-specific insights and recommendations
        """
        # Section heading
        services_heading = doc.add_heading("Service Resources", level=1)
        services_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Group resources by service
        services = defaultdict(list)
        for resource in bom_data.resources:
            service = resource.get("service", "Unknown")
            services[service].append(resource)
            
        # Create section for each service
        for service, resources in sorted(services.items()):
            self._add_single_service_section(doc, service, resources)
            
    def _add_single_service_section(self, doc: Document, service: str, resources: List[Dict]):
        """Add section for a specific service."""
        # Service heading
        service_heading = doc.add_heading(f"{service} Resources", level=2)
        service_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Service summary
        compliant_count = sum(1 for r in resources if r.get("compliance_status") == "compliant")
        non_compliant_count = len(resources) - compliant_count
        
        summary_para = doc.add_paragraph()
        summary_text = (
            f"Total {service} resources: {len(resources)} "
            f"(Compliant: {compliant_count}, Non-compliant: {non_compliant_count})"
        )
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.size = self.styles["body_font_size"]
        summary_run.font.bold = True
        
        # Resource table
        if resources:
            # Get common fields for this service
            common_fields = self._get_common_fields(resources)
            
            # Create table
            resource_table = doc.add_table(rows=len(resources) + 1, cols=len(common_fields))
            resource_table.style = 'Table Grid'
            
            # Headers
            header_row = resource_table.rows[0]
            for i, field in enumerate(common_fields):
                header_row.cells[i].text = field.replace("_", " ").title()
                
            # Format headers
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = self.styles["table_font_size"]
                        run.font.color.rgb = self.styles["primary_color"]
                        
            # Resource data
            for i, resource in enumerate(resources, 1):
                row = resource_table.rows[i]
                for j, field in enumerate(common_fields):
                    value = resource.get(field, "")
                    if isinstance(value, dict):
                        value = str(value)
                    elif isinstance(value, list):
                        value = ", ".join(str(item) for item in value)
                    
                    row.cells[j].text = str(value) if value is not None else ""
                    
                    # Format cells
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = self.styles["table_font_size"]
                            
                            # Color code compliance status
                            if field == "compliance_status":
                                if value == "compliant":
                                    run.font.color.rgb = self.styles["accent_color"]
                                elif value == "non_compliant":
                                    run.font.color.rgb = self.styles["danger_color"]
                                    
        doc.add_paragraph()  # Add spacing
        
    def _add_network_analysis_section(self, doc: Document, bom_data: BOMData):
        """
        Add network analysis section with CIDR utilization details.
        
        Enhanced network analysis with:
        - CIDR utilization details and capacity planning
        - VPC and subnet analysis
        - Network resource mapping
        """
        # Section heading
        network_heading = doc.add_heading("Network Analysis", level=1)
        network_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        network_analysis = bom_data.network_analysis
        
        # Network overview
        doc.add_heading("Network Overview", level=2)
        
        overview_para = doc.add_paragraph()
        overview_text = (
            f"The infrastructure spans {network_analysis.get('total_vpcs', 0)} VPCs "
            f"with {network_analysis.get('total_subnets', 0)} subnets. "
            f"This section provides detailed analysis of network utilization and capacity planning."
        )
        overview_run = overview_para.add_run(overview_text)
        overview_run.font.size = self.styles["body_font_size"]
        
        # VPC utilization analysis
        vpc_utilization = network_analysis.get("vpc_utilization", {})
        if vpc_utilization:
            doc.add_heading("VPC Utilization Analysis", level=2)
            
            # VPC table
            vpc_table = doc.add_table(rows=len(vpc_utilization) + 1, cols=5)
            vpc_table.style = 'Table Grid'
            
            # Headers
            headers = ["VPC ID", "VPC Name", "CIDR Block", "Utilization %", "Available IPs"]
            header_row = vpc_table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                
            # Format headers
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = self.styles["table_font_size"]
                        run.font.color.rgb = self.styles["primary_color"]
                        
            # VPC data
            for i, (vpc_id, vpc_data) in enumerate(vpc_utilization.items(), 1):
                row = vpc_table.rows[i]
                
                row.cells[0].text = vpc_id
                row.cells[1].text = vpc_data.get("name", "")
                row.cells[2].text = vpc_data.get("cidr_block", "")
                row.cells[3].text = f"{vpc_data.get('utilization_percentage', 0):.1f}%"
                row.cells[4].text = str(vpc_data.get("available_ips", 0))
                
                # Format cells
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = self.styles["table_font_size"]
                            
                # Color code utilization
                utilization = vpc_data.get("utilization_percentage", 0)
                utilization_cell = row.cells[3]
                for paragraph in utilization_cell.paragraphs:
                    for run in paragraph.runs:
                        if utilization > 80:
                            run.font.color.rgb = self.styles["danger_color"]
                        elif utilization > 60:
                            run.font.color.rgb = self.styles["warning_color"]
                        else:
                            run.font.color.rgb = self.styles["accent_color"]
                            
        # Network recommendations
        doc.add_heading("Network Recommendations", level=2)
        
        network_recommendations = [
            "Monitor VPC utilization and plan for capacity expansion",
            "Review subnet allocation for optimal resource distribution",
            "Consider VPC peering for cross-VPC communication needs",
            "Implement network segmentation best practices",
            "Regular review of security group rules and NACLs"
        ]
        
        for recommendation in network_recommendations:
            rec_para = doc.add_paragraph()
            rec_run = rec_para.add_run(f"• {recommendation}")
            rec_run.font.size = self.styles["body_font_size"]
            
        doc.add_page_break()
        
    def _add_security_analysis_section(self, doc: Document, bom_data: BOMData):
        """
        Add security analysis section with risk assessment summaries and recommendations.
        
        Enhanced security analysis with:
        - Risk assessment summaries and recommendations
        - Security group analysis
        - Security findings and remediation guidance
        """
        # Section heading
        security_heading = doc.add_heading("Security Analysis", level=1)
        security_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        security_analysis = bom_data.security_analysis
        
        # Security overview
        doc.add_heading("Security Overview", level=2)
        
        overview_para = doc.add_paragraph()
        overview_text = (
            f"Security analysis covers {security_analysis.get('total_security_groups', 0)} security groups "
            f"with {security_analysis.get('high_risk_rules', 0)} high-risk rules identified. "
            f"This section provides detailed risk assessment and remediation recommendations."
        )
        overview_run = overview_para.add_run(overview_text)
        overview_run.font.size = self.styles["body_font_size"]
        
        # High-risk rules analysis
        overly_permissive_rules = security_analysis.get("overly_permissive_rules", [])
        if overly_permissive_rules:
            doc.add_heading("High-Risk Security Rules", level=2)
            
            # Risk table
            risk_table = doc.add_table(rows=len(overly_permissive_rules) + 1, cols=4)
            risk_table.style = 'Table Grid'
            
            # Headers
            headers = ["Security Group", "Rule", "Risk Level", "Recommendation"]
            header_row = risk_table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                
            # Format headers
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = self.styles["table_font_size"]
                        run.font.color.rgb = self.styles["primary_color"]
                        
            # Risk data
            for i, rule in enumerate(overly_permissive_rules, 1):
                row = risk_table.rows[i]
                
                row.cells[0].text = rule.get("group_id", "")
                row.cells[1].text = rule.get("rule", "")
                row.cells[2].text = rule.get("risk_level", "")
                row.cells[3].text = "Restrict source to specific CIDR blocks"
                
                # Format cells
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = self.styles["table_font_size"]
                            
                # Color code risk level
                risk_level = rule.get("risk_level", "")
                risk_cell = row.cells[2]
                for paragraph in risk_cell.paragraphs:
                    for run in paragraph.runs:
                        if risk_level.lower() == "high":
                            run.font.color.rgb = self.styles["danger_color"]
                        elif risk_level.lower() == "medium":
                            run.font.color.rgb = self.styles["warning_color"]
                        else:
                            run.font.color.rgb = self.styles["accent_color"]
                            
        # Security recommendations
        doc.add_heading("Security Recommendations", level=2)
        
        security_recommendations = [
            "Review and restrict overly permissive security group rules",
            "Implement least privilege access principles",
            "Regular security group audits and cleanup",
            "Use AWS Config rules for continuous compliance monitoring",
            "Enable VPC Flow Logs for network traffic analysis",
            "Implement AWS GuardDuty for threat detection",
            "Regular penetration testing and vulnerability assessments"
        ]
        
        for recommendation in security_recommendations:
            rec_para = doc.add_paragraph()
            rec_run = rec_para.add_run(f"• {recommendation}")
            rec_run.font.size = self.styles["body_font_size"]
            
        doc.add_page_break()
        
    def _add_compliance_details_section(self, doc: Document, bom_data: BOMData):
        """Add detailed compliance information section."""
        # Section heading
        compliance_heading = doc.add_heading("Compliance Details", level=1)
        compliance_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Compliance breakdown by service
        doc.add_heading("Compliance Breakdown by Service", level=2)
        
        # Count compliance by service
        service_compliance = defaultdict(lambda: {"compliant": 0, "non_compliant": 0})
        
        for resource in bom_data.resources:
            service = resource.get("service", "Unknown")
            status = resource.get("compliance_status", "unknown")
            
            if status == "compliant":
                service_compliance[service]["compliant"] += 1
            elif status == "non_compliant":
                service_compliance[service]["non_compliant"] += 1
                
        # Service compliance table
        compliance_table = doc.add_table(rows=len(service_compliance) + 1, cols=5)
        compliance_table.style = 'Table Grid'
        
        # Headers
        headers = ["Service", "Compliant", "Non-Compliant", "Total", "Compliance %"]
        header_row = compliance_table.rows[0]
        for i, header in enumerate(headers):
            header_row.cells[i].text = header
            
        # Format headers
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = self.styles["table_font_size"]
                    run.font.color.rgb = self.styles["primary_color"]
                    
        # Service data
        for i, (service, counts) in enumerate(sorted(service_compliance.items()), 1):
            row = compliance_table.rows[i]
            
            compliant = counts["compliant"]
            non_compliant = counts["non_compliant"]
            total = compliant + non_compliant
            compliance_pct = (compliant / total * 100) if total > 0 else 0
            
            row.cells[0].text = service
            row.cells[1].text = str(compliant)
            row.cells[2].text = str(non_compliant)
            row.cells[3].text = str(total)
            row.cells[4].text = f"{compliance_pct:.1f}%"
            
            # Format cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = self.styles["table_font_size"]
                        
            # Color code compliance values
            compliant_cell = row.cells[1]
            for paragraph in compliant_cell.paragraphs:
                for run in paragraph.runs:
                    if compliant > 0:
                        run.font.color.rgb = self.styles["accent_color"]
                        
            non_compliant_cell = row.cells[2]
            for paragraph in non_compliant_cell.paragraphs:
                for run in paragraph.runs:
                    if non_compliant > 0:
                        run.font.color.rgb = self.styles["danger_color"]
                        
        # Non-compliant resources details
        non_compliant_resources = [
            r for r in bom_data.resources 
            if r.get("compliance_status") == "non_compliant"
        ]
        
        if non_compliant_resources:
            doc.add_heading("Non-Compliant Resources", level=2)
            
            # Non-compliant table
            nc_table = doc.add_table(rows=len(non_compliant_resources) + 1, cols=6)
            nc_table.style = 'Table Grid'
            
            # Headers
            headers = ["Service", "Type", "ID", "Name", "Region", "Issues"]
            header_row = nc_table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                
            # Format headers
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = self.styles["table_font_size"]
                        run.font.color.rgb = self.styles["primary_color"]
                        
            # Resource data
            for i, resource in enumerate(non_compliant_resources, 1):
                row = nc_table.rows[i]
                
                row.cells[0].text = resource.get("service", "")
                row.cells[1].text = resource.get("type", "")
                row.cells[2].text = resource.get("id", "")
                row.cells[3].text = resource.get("name", "")
                row.cells[4].text = resource.get("region", "")
                row.cells[5].text = "Missing required tags"
                
                # Format cells
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = self.styles["table_font_size"]
                            
        doc.add_page_break()
        
    def _add_appendices(self, doc: Document, bom_data: BOMData):
        """Add appendices with additional information."""
        # Section heading
        appendix_heading = doc.add_heading("Appendices", level=1)
        appendix_heading.runs[0].font.color.rgb = self.styles["primary_color"]
        
        # Appendix A: Generation metadata
        doc.add_heading("Appendix A: Report Generation Details", level=2)
        
        metadata = bom_data.generation_metadata
        
        metadata_para = doc.add_paragraph()
        metadata_text = (
            f"This report was generated on {metadata.get('generated_at', 'Unknown')} "
            f"and processed {metadata.get('total_resources', 0)} resources "
            f"in {metadata.get('processing_time_seconds', 0):.2f} seconds. "
            f"The data source was {metadata.get('data_source', 'AWS Resource Inventory')}."
        )
        metadata_run = metadata_para.add_run(metadata_text)
        metadata_run.font.size = self.styles["body_font_size"]
        
        # Appendix B: Custom attributes
        if bom_data.custom_attributes:
            doc.add_heading("Appendix B: Custom Attributes", level=2)
            
            custom_para = doc.add_paragraph()
            custom_text = (
                f"The following custom attributes were configured for this report: "
                f"{', '.join(bom_data.custom_attributes)}. "
                f"These attributes provide organization-specific metadata for resources."
            )
            custom_run = custom_para.add_run(custom_text)
            custom_run.font.size = self.styles["body_font_size"]
            
    def _apply_document_branding(self, doc: Document):
        """Apply branding to the document."""
        # Set document properties
        core_props = doc.core_properties
        core_props.title = f"{self.config.branding.company_name} - Cloud BOM Report"
        core_props.author = "InvenTag Cloud BOM Generator"
        core_props.subject = "AWS Resource Inventory and Compliance Report"
        core_props.comments = "Professional AWS resource inventory and compliance report"
        
    def _get_common_fields(self, resources: List[Dict]) -> List[str]:
        """Get common fields across resources for table generation."""
        if not resources:
            return []
            
        # Priority fields to show first
        priority_fields = [
            "id", "name", "type", "region", "compliance_status", 
            "account_id", "vpc_id", "subnet_id"
        ]
        
        # Get all fields from all resources
        all_fields = set()
        for resource in resources:
            all_fields.update(resource.keys())
            
        # Filter out complex fields
        simple_fields = []
        for field in all_fields:
            # Check if field has simple values
            sample_value = resources[0].get(field)
            if not isinstance(sample_value, (dict, list)) or field in priority_fields:
                simple_fields.append(field)
                
        # Sort with priority fields first
        sorted_fields = []
        for field in priority_fields:
            if field in simple_fields:
                sorted_fields.append(field)
                simple_fields.remove(field)
                
        # Add remaining fields
        sorted_fields.extend(sorted(simple_fields))
        
        # Limit to reasonable number of columns
        return sorted_fields[:8]
        
    def _generate_key_findings(self, bom_data: BOMData) -> List[str]:
        """Generate key findings based on BOM data."""
        findings = []
        
        compliance = bom_data.compliance_summary
        total_resources = compliance.get("total_resources", 0)
        compliant_resources = compliance.get("compliant_resources", 0)
        
        if total_resources > 0:
            compliance_pct = (compliant_resources / total_resources) * 100
            
            if compliance_pct >= 90:
                findings.append(f"Excellent compliance rate of {compliance_pct:.1f}% demonstrates strong governance practices")
            elif compliance_pct >= 70:
                findings.append(f"Good compliance rate of {compliance_pct:.1f}% with opportunities for improvement")
            else:
                findings.append(f"Compliance rate of {compliance_pct:.1f}% requires immediate attention and remediation")
                
        # Network findings
        network_analysis = bom_data.network_analysis
        if network_analysis:
            total_vpcs = network_analysis.get("total_vpcs", 0)
            if total_vpcs > 0:
                findings.append(f"Infrastructure spans {total_vpcs} VPCs requiring network optimization review")
                
        # Security findings
        security_analysis = bom_data.security_analysis
        if security_analysis:
            high_risk_rules = security_analysis.get("high_risk_rules", 0)
            if high_risk_rules > 0:
                findings.append(f"Identified {high_risk_rules} high-risk security rules requiring immediate remediation")
            else:
                findings.append("No high-risk security rules identified - security posture is strong")
                
        # Service diversity
        services = set(resource.get("service", "Unknown") for resource in bom_data.resources)
        findings.append(f"Infrastructure utilizes {len(services)} different AWS services")
        
        return findings
        
    def apply_branding(self, document: Document) -> Document:
        """Apply branding configuration to Word document."""
        self._apply_document_branding(document)
        return document