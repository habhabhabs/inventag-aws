#!/usr/bin/env python3
"""
Unit tests for WordDocumentBuilder.

Tests the Word document generation functionality with professional formatting.
"""

import unittest
import tempfile
import os
import shutil
from unittest.mock import Mock, patch, MagicMock
from datetime import datetime, timezone

# Import the modules to test
from inventag.reporting.word_builder import WordDocumentBuilder
from inventag.reporting.document_generator import DocumentConfig, BrandingConfig, DocumentGenerationResult
from inventag.reporting.bom_processor import BOMData

# Check if python-docx is available
try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class TestWordDocumentBuilder(unittest.TestCase):
    """Test cases for WordDocumentBuilder."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        
        # Create test BOM data
        self.test_bom_data = BOMData(
            resources=[
                {
                    "service": "EC2",
                    "type": "Instance",
                    "id": "i-1234567890abcdef0",
                    "name": "web-server-1",
                    "region": "us-east-1",
                    "compliance_status": "compliant",
                    "tags": {
                        "Name": "web-server-1",
                        "Environment": "production"
                    },
                    "vpc_id": "vpc-12345678",
                    "instance_type": "t3.medium"
                },
                {
                    "service": "S3",
                    "type": "Bucket",
                    "id": "test-bucket",
                    "name": "test-bucket",
                    "region": "us-east-1",
                    "compliance_status": "non_compliant",
                    "tags": {
                        "Name": "test-bucket"
                    },
                    "encryption": "AES256"
                },
                {
                    "service": "EC2",
                    "type": "Volume",
                    "id": "vol-1234567890abcdef0",
                    "name": "web-server-volume",
                    "region": "us-east-1",
                    "compliance_status": "compliant",
                    "vpc_id": "vpc-12345678",
                    "size": 20
                }
            ],
            network_analysis={
                "total_vpcs": 1,
                "total_subnets": 2,
                "vpc_utilization": {
                    "vpc-12345678": {
                        "name": "main-vpc",
                        "cidr_block": "10.0.0.0/16",
                        "utilization_percentage": 25.5,
                        "available_ips": 65280
                    }
                }
            },
            security_analysis={
                "total_security_groups": 3,
                "high_risk_rules": 1,
                "overly_permissive_rules": [
                    {
                        "group_id": "sg-12345678",
                        "rule": "0.0.0.0/0:22",
                        "risk_level": "high"
                    }
                ]
            },
            compliance_summary={
                "total_resources": 3,
                "compliant_resources": 2,
                "non_compliant_resources": 1,
                "compliance_percentage": 66.7
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 3,
                "processing_time_seconds": 1.5,
                "data_source": "aws_resource_inventory.py"
            },
            custom_attributes=["inventag:remarks", "inventag:costcenter"]
        )
        
        # Create test configuration
        self.test_config = DocumentConfig(
            output_directory=self.temp_dir,
            branding=BrandingConfig(
                company_name="Test Corporation",
                color_scheme={
                    "primary": "366092",
                    "accent": "70AD47",
                    "danger": "C5504B",
                    "warning": "FFC000"
                },
                font_family="Calibri"
            )
        )
        
    def tearDown(self):
        """Clean up test fixtures."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        
    def test_word_builder_initialization(self):
        """Test WordDocumentBuilder initialization."""
        builder = WordDocumentBuilder(self.test_config)
        
        self.assertIsInstance(builder, WordDocumentBuilder)
        self.assertEqual(builder.config, self.test_config)
        self.assertIsInstance(builder.styles, dict)
        
    def test_can_handle_format(self):
        """Test format handling capability."""
        builder = WordDocumentBuilder(self.test_config)
        
        self.assertTrue(builder.can_handle_format("word"))
        self.assertTrue(builder.can_handle_format("WORD"))
        self.assertFalse(builder.can_handle_format("excel"))
        self.assertFalse(builder.can_handle_format("csv"))
        
    def test_validate_dependencies(self):
        """Test dependency validation."""
        builder = WordDocumentBuilder(self.test_config)
        dependencies = builder.validate_dependencies()
        
        if PYTHON_DOCX_AVAILABLE:
            self.assertEqual(dependencies, [])
        else:
            self.assertIn("python-docx library not available", dependencies[0])
            
    def test_styles_initialization(self):
        """Test Word styles initialization."""
        builder = WordDocumentBuilder(self.test_config)
        
        if PYTHON_DOCX_AVAILABLE:
            self.assertIn("title_font_size", builder.styles)
            self.assertIn("heading1_font_size", builder.styles)
            self.assertIn("body_font_size", builder.styles)
            self.assertIn("primary_color", builder.styles)
            self.assertIn("font_name", builder.styles)
            
    def test_hex_to_rgb_conversion(self):
        """Test hex color to RGB conversion."""
        builder = WordDocumentBuilder(self.test_config)
        
        if PYTHON_DOCX_AVAILABLE:
            # Test normal 6-character hex
            rgb = builder._hex_to_rgb("FF0000")
            self.assertEqual(rgb, docx.shared.RGBColor(255, 0, 0))
            
            # Test 3-character hex
            rgb = builder._hex_to_rgb("F00")
            self.assertEqual(rgb, docx.shared.RGBColor(255, 0, 0))
            
            # Test with # prefix
            rgb = builder._hex_to_rgb("#00FF00")
            self.assertEqual(rgb, docx.shared.RGBColor(0, 255, 0))
            
            # Test invalid color (should default to black)
            rgb = builder._hex_to_rgb("invalid")
            self.assertEqual(rgb, docx.shared.RGBColor(0, 0, 0))
            
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_generate_document_success(self):
        """Test successful Word document generation."""
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "test_report.docx")
        
        result = builder.generate_document(self.test_bom_data, output_path)
        
        self.assertIsInstance(result, DocumentGenerationResult)
        self.assertTrue(result.success)
        self.assertEqual(result.format_type, "word")
        self.assertEqual(result.filename, "test_report.docx")
        self.assertIsNone(result.error_message)
        self.assertGreater(result.generation_time_seconds, 0)
        
        # Verify file was created
        self.assertTrue(os.path.exists(output_path))
        self.assertGreater(os.path.getsize(output_path), 0)
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_document_structure(self):
        """Test Word document structure and content."""
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "structure_test.docx")
        
        result = builder.generate_document(self.test_bom_data, output_path)
        self.assertTrue(result.success)
        
        # Load and verify document structure
        from docx import Document
        doc = Document(output_path)
        
        # Check that document has content
        self.assertGreater(len(doc.paragraphs), 10)
        
        # Check for expected sections (by looking for headings)
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text)
                
        expected_headings = [
            "Table of Contents",
            "Executive Summary",
            "Service Resources",
            "Network Analysis",
            "Security Analysis",
            "Compliance Details"
        ]
        
        # Also check all paragraph text for title
        all_text = [p.text for p in doc.paragraphs]
        title_found = any("Test Corporation" in text for text in all_text)
        self.assertTrue(title_found, "Company title not found in document")
        
        for expected in expected_headings:
            found = any(expected in heading for heading in headings)
            self.assertTrue(found, f"Expected heading '{expected}' not found in document")
            
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_document_tables(self):
        """Test Word document table generation."""
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "tables_test.docx")
        
        result = builder.generate_document(self.test_bom_data, output_path)
        self.assertTrue(result.success)
        
        # Load and verify document tables
        from docx import Document
        doc = Document(output_path)
        
        # Should have multiple tables
        self.assertGreater(len(doc.tables), 3)
        
        # Check that tables have content
        for table in doc.tables:
            self.assertGreater(len(table.rows), 0)
            self.assertGreater(len(table.columns), 0)
            
    def test_generate_document_without_python_docx(self):
        """Test document generation when python-docx is not available."""
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "no_docx_test.docx")
        
        # Mock python-docx as not available
        with patch('inventag.reporting.word_builder.PYTHON_DOCX_AVAILABLE', False):
            result = builder.generate_document(self.test_bom_data, output_path)
            
        self.assertFalse(result.success)
        self.assertEqual(result.format_type, "word")
        self.assertIn("python-docx library not available", result.error_message)
        
    def test_generate_document_with_exception(self):
        """Test document generation with exception handling."""
        builder = WordDocumentBuilder(self.test_config)
        
        # Create invalid BOM data to trigger an exception during processing
        invalid_bom_data = BOMData(
            resources=[{"invalid": "data"}],  # Missing required fields
            generation_metadata=None  # This should cause an error
        )
        
        output_path = os.path.join(self.temp_dir, "exception_test.docx")
        result = builder.generate_document(invalid_bom_data, output_path)
        
        # Should handle the exception gracefully
        self.assertFalse(result.success)
        self.assertEqual(result.format_type, "word")
        self.assertIsNotNone(result.error_message)
        
    def test_get_common_fields_functionality(self):
        """Test common fields extraction functionality."""
        builder = WordDocumentBuilder(self.test_config)
        
        test_resources = [
            {
                "id": "resource-1",
                "name": "Resource 1",
                "type": "Instance",
                "region": "us-east-1",
                "compliance_status": "compliant",
                "complex_field": {"nested": "data"},
                "list_field": ["item1", "item2"]
            },
            {
                "id": "resource-2",
                "name": "Resource 2",
                "type": "Volume",
                "region": "us-west-2",
                "compliance_status": "non_compliant",
                "size": 100
            }
        ]
        
        common_fields = builder._get_common_fields(test_resources)
        
        # Should include priority fields first
        self.assertIn("id", common_fields)
        self.assertIn("name", common_fields)
        self.assertIn("type", common_fields)
        self.assertIn("region", common_fields)
        self.assertIn("compliance_status", common_fields)
        
        # Should limit to reasonable number of columns
        self.assertLessEqual(len(common_fields), 8)
        
    def test_generate_key_findings(self):
        """Test key findings generation."""
        builder = WordDocumentBuilder(self.test_config)
        
        findings = builder._generate_key_findings(self.test_bom_data)
        
        self.assertIsInstance(findings, list)
        self.assertGreater(len(findings), 0)
        
        # Should include compliance finding
        compliance_finding = any("compliance rate" in finding.lower() for finding in findings)
        self.assertTrue(compliance_finding, "Compliance finding not generated")
        
        # Should include security finding (we have high risk rules)
        security_finding = any("high-risk security rules" in finding.lower() for finding in findings)
        self.assertTrue(security_finding, "Security finding not generated")
        
        # Should include service diversity finding
        service_finding = any("aws services" in finding.lower() for finding in findings)
        self.assertTrue(service_finding, "Service diversity finding not generated")
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_branding_application(self):
        """Test branding application to document."""
        builder = WordDocumentBuilder(self.test_config)
        
        from docx import Document
        doc = Document()
        
        branded_doc = builder.apply_branding(doc)
        
        self.assertEqual(branded_doc.core_properties.title, "Test Corporation - Cloud BOM Report")
        self.assertEqual(branded_doc.core_properties.author, "InvenTag Cloud BOM Generator")
        
    def test_custom_branding_colors(self):
        """Test custom branding colors in styles."""
        custom_config = DocumentConfig(
            branding=BrandingConfig(
                color_scheme={
                    "primary": "FF0000",
                    "accent": "00FF00",
                    "danger": "0000FF",
                    "warning": "FFFF00"
                },
                font_family="Arial"
            )
        )
        
        builder = WordDocumentBuilder(custom_config)
        
        if PYTHON_DOCX_AVAILABLE:
            # Check that custom colors are used in styles
            self.assertEqual(builder.styles["primary_color"], docx.shared.RGBColor(255, 0, 0))
            self.assertEqual(builder.styles["accent_color"], docx.shared.RGBColor(0, 255, 0))
            self.assertEqual(builder.styles["font_name"], "Arial")
            
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_empty_data_handling(self):
        """Test handling of empty or minimal data."""
        empty_bom_data = BOMData(
            resources=[],
            network_analysis={},
            security_analysis={},
            compliance_summary={},
            generation_metadata={}
        )
        
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "empty_test.docx")
        
        result = builder.generate_document(empty_bom_data, output_path)
        
        # Should still succeed but with minimal content
        self.assertTrue(result.success)
        self.assertTrue(os.path.exists(output_path))
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_large_dataset_handling(self):
        """Test handling of larger datasets."""
        # Create larger test dataset
        large_resources = []
        for i in range(50):
            large_resources.append({
                "service": f"Service{i % 5}",
                "type": "Resource",
                "id": f"resource-{i:03d}",
                "name": f"Resource {i}",
                "region": "us-east-1",
                "compliance_status": "compliant" if i % 2 == 0 else "non_compliant"
            })
            
        large_bom_data = BOMData(
            resources=large_resources,
            network_analysis={"total_vpcs": 10, "total_subnets": 50},
            security_analysis={"total_security_groups": 25, "high_risk_rules": 5},
            compliance_summary={
                "total_resources": 50,
                "compliant_resources": 25,
                "non_compliant_resources": 25,
                "compliance_percentage": 50.0
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 50
            }
        )
        
        builder = WordDocumentBuilder(self.test_config)
        output_path = os.path.join(self.temp_dir, "large_test.docx")
        
        result = builder.generate_document(large_bom_data, output_path)
        
        self.assertTrue(result.success)
        self.assertTrue(os.path.exists(output_path))
        self.assertGreater(os.path.getsize(output_path), 20000)  # Should be reasonably large file
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_service_section_generation(self):
        """Test service-specific section generation."""
        builder = WordDocumentBuilder(self.test_config)
        
        # Test with resources from multiple services
        multi_service_resources = [
            {"service": "EC2", "id": "i-123", "compliance_status": "compliant"},
            {"service": "EC2", "id": "i-456", "compliance_status": "non_compliant"},
            {"service": "S3", "id": "bucket-1", "compliance_status": "compliant"},
            {"service": "RDS", "id": "db-1", "compliance_status": "compliant"}
        ]
        
        multi_service_bom_data = BOMData(
            resources=multi_service_resources,
            network_analysis={},
            security_analysis={},
            compliance_summary={
                "total_resources": 4,
                "compliant_resources": 3,
                "non_compliant_resources": 1,
                "compliance_percentage": 75.0
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 4
            }
        )
        
        output_path = os.path.join(self.temp_dir, "multi_service_test.docx")
        result = builder.generate_document(multi_service_bom_data, output_path)
        
        self.assertTrue(result.success)
        
        # Load and check for service sections
        from docx import Document
        doc = Document(output_path)
        
        # Should have sections for each service
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # Check for service-specific headings
        ec2_found = any("EC2" in heading for heading in headings)
        s3_found = any("S3" in heading for heading in headings)
        rds_found = any("RDS" in heading for heading in headings)
        
        self.assertTrue(ec2_found, "EC2 section not found")
        self.assertTrue(s3_found, "S3 section not found")
        self.assertTrue(rds_found, "RDS section not found")


if __name__ == '__main__':
    unittest.main()