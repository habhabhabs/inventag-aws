#!/usr/bin/env python3
"""
Integration tests for Word document generation.

Tests the complete Word document generation workflow.
"""

import unittest
import tempfile
import os
import shutil
from datetime import datetime, timezone

from inventag.reporting.document_generator import DocumentGenerator, DocumentConfig, BrandingConfig
from inventag.reporting.bom_processor import BOMData

# Check if python-docx is available
try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class TestWordIntegration(unittest.TestCase):
    """Integration test cases for Word document generation."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        
        # Create comprehensive test BOM data
        self.test_bom_data = BOMData(
            resources=[
                {
                    "service": "EC2",
                    "type": "Instance",
                    "id": "i-1234567890abcdef0",
                    "name": "web-server-1",
                    "region": "us-east-1",
                    "account_id": "123456789012",
                    "arn": "arn:aws:ec2:us-east-1:123456789012:instance/i-1234567890abcdef0",
                    "compliance_status": "compliant",
                    "tags": {
                        "Name": "web-server-1",
                        "Environment": "production",
                        "inventag:remarks": "Primary web server",
                        "inventag:costcenter": "IT-001"
                    },
                    "vpc_id": "vpc-12345678",
                    "subnet_id": "subnet-12345678",
                    "security_groups": ["sg-12345678"],
                    "instance_type": "t3.medium",
                    "state": "running"
                },
                {
                    "service": "S3",
                    "type": "Bucket",
                    "id": "test-bucket-12345",
                    "name": "test-bucket-12345",
                    "region": "us-east-1",
                    "account_id": "123456789012",
                    "arn": "arn:aws:s3:::test-bucket-12345",
                    "compliance_status": "non_compliant",
                    "tags": {
                        "Name": "test-bucket-12345",
                        "Environment": "development"
                    },
                    "encryption": "AES256",
                    "versioning": "Enabled",
                    "public_access_blocked": True
                },
                {
                    "service": "RDS",
                    "type": "DBInstance",
                    "id": "database-1",
                    "name": "database-1",
                    "region": "us-west-2",
                    "account_id": "123456789012",
                    "arn": "arn:aws:rds:us-west-2:123456789012:db:database-1",
                    "compliance_status": "compliant",
                    "tags": {
                        "Name": "database-1",
                        "Environment": "production",
                        "inventag:remarks": "Main application database",
                        "inventag:costcenter": "IT-001"
                    },
                    "engine": "mysql",
                    "engine_version": "8.0.35",
                    "instance_class": "db.t3.micro",
                    "allocated_storage": 20
                }
            ],
            network_analysis={
                "total_vpcs": 2,
                "total_subnets": 6,
                "vpc_utilization": {
                    "vpc-12345678": {
                        "name": "main-vpc",
                        "cidr_block": "10.0.0.0/16",
                        "utilization_percentage": 25.5,
                        "available_ips": 65280
                    },
                    "vpc-87654321": {
                        "name": "dev-vpc",
                        "cidr_block": "10.1.0.0/16",
                        "utilization_percentage": 15.2,
                        "available_ips": 65400
                    }
                }
            },
            security_analysis={
                "total_security_groups": 8,
                "high_risk_rules": 2,
                "overly_permissive_rules": [
                    {
                        "group_id": "sg-12345678",
                        "rule": "0.0.0.0/0:22",
                        "risk_level": "high"
                    },
                    {
                        "group_id": "sg-87654321",
                        "rule": "0.0.0.0/0:3389",
                        "risk_level": "high"
                    }
                ]
            },
            compliance_summary={
                "total_resources": 3,
                "compliant_resources": 2,
                "non_compliant_resources": 1,
                "compliance_percentage": 66.7
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 3,
                "processing_time_seconds": 2.1,
                "data_source": "aws_resource_inventory.py"
            },
            custom_attributes=["inventag:remarks", "inventag:costcenter"]
        )
        
    def tearDown(self):
        """Clean up test fixtures."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_document_generation_integration(self):
        """Test complete Word document generation workflow."""
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            filename_template="integration_test_{timestamp}",
            branding=BrandingConfig(
                company_name="Integration Test Corp",
                color_scheme={
                    "primary": "2E75B6",
                    "accent": "70AD47",
                    "danger": "C5504B",
                    "warning": "FFC000"
                },
                font_family="Calibri"
            ),
            validate_before_generation=True
        )
        
        generator = DocumentGenerator(config)
        
        # Generate documents
        summary = generator.generate_bom_documents(self.test_bom_data)
        
        # Verify generation summary
        self.assertEqual(summary.total_formats, 1)
        self.assertEqual(summary.successful_formats, 1)
        self.assertEqual(summary.failed_formats, 0)
        self.assertEqual(len(summary.results), 1)
        
        result = summary.results[0]
        self.assertTrue(result.success)
        self.assertEqual(result.format_type, "word")
        self.assertTrue(result.filename.endswith(".docx"))
        self.assertIsNone(result.error_message)
        self.assertGreater(result.generation_time_seconds, 0)
        
        # Verify file was created
        output_file = os.path.join(self.temp_dir, result.filename)
        self.assertTrue(os.path.exists(output_file))
        self.assertGreater(os.path.getsize(output_file), 10000)  # Should be substantial file
        
        # Load and verify Word content
        doc = docx.Document(output_file)
        
        # Verify document has substantial content
        self.assertGreater(len(doc.paragraphs), 20)
        self.assertGreater(len(doc.tables), 3)
        
        # Verify expected sections exist (by looking for headings)
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text)
                
        expected_headings = [
            "Table of Contents",
            "Executive Summary",
            "Service Resources",
            "Network Analysis",
            "Security Analysis",
            "Compliance Details"
        ]
        
        for expected in expected_headings:
            found = any(expected in heading for heading in headings)
            self.assertTrue(found, f"Expected heading '{expected}' not found in document")
            
        # Verify company name appears in document
        all_text = [p.text for p in doc.paragraphs]
        company_found = any("Integration Test Corp" in text for text in all_text)
        self.assertTrue(company_found, "Company name not found in document")
        
        # Verify document properties
        self.assertEqual(doc.core_properties.title, "Integration Test Corp - Cloud BOM Report")
        self.assertEqual(doc.core_properties.author, "InvenTag Cloud BOM Generator")
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_with_custom_branding(self):
        """Test Word generation with custom branding."""
        custom_branding = BrandingConfig(
            company_name="Custom Branding Test",
            color_scheme={
                "primary": "FF6B35",
                "accent": "004E89",
                "danger": "D62828",
                "warning": "F77F00"
            },
            font_family="Arial"
        )
        
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            branding=custom_branding,
            filename_template="branded_test_{timestamp}"
        )
        
        generator = DocumentGenerator(config)
        summary = generator.generate_bom_documents(self.test_bom_data)
        
        # Should succeed
        self.assertEqual(summary.successful_formats, 1)
        result = summary.results[0]
        self.assertTrue(result.success)
        
        # Verify file and branding
        output_file = os.path.join(self.temp_dir, result.filename)
        doc = docx.Document(output_file)
        
        # Check document properties
        self.assertEqual(doc.core_properties.title, "Custom Branding Test - Cloud BOM Report")
        self.assertEqual(doc.core_properties.author, "InvenTag Cloud BOM Generator")
        
        # Verify custom company name appears
        all_text = [p.text for p in doc.paragraphs]
        company_found = any("Custom Branding Test" in text for text in all_text)
        self.assertTrue(company_found, "Custom company name not found in document")
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_with_large_dataset(self):
        """Test Word generation with larger dataset."""
        # Create larger test dataset
        large_resources = []
        services = ["EC2", "S3", "RDS", "LAMBDA", "VPC", "ELB", "CLOUDFORMATION"]
        
        for i in range(30):  # Smaller than Excel test to keep Word doc manageable
            service = services[i % len(services)]
            large_resources.append({
                "service": service,
                "type": f"Resource{i % 3}",
                "id": f"resource-{i:03d}",
                "name": f"Resource {i}",
                "region": "us-east-1" if i % 2 == 0 else "us-west-2",
                "account_id": "123456789012",
                "compliance_status": "compliant" if i % 3 == 0 else "non_compliant",
                "tags": {
                    "Name": f"Resource {i}",
                    "Environment": "production" if i % 2 == 0 else "development",
                    "inventag:costcenter": f"CC-{i % 5:03d}"
                }
            })
            
        large_bom_data = BOMData(
            resources=large_resources,
            network_analysis={
                "total_vpcs": 3,
                "total_subnets": 12,
                "vpc_utilization": {
                    f"vpc-{i:08d}": {
                        "name": f"vpc-{i}",
                        "cidr_block": f"10.{i}.0.0/16",
                        "utilization_percentage": (i * 15) % 100,
                        "available_ips": 65536 - (i * 100)
                    }
                    for i in range(3)
                }
            },
            security_analysis={
                "total_security_groups": 10,
                "high_risk_rules": 2,
                "overly_permissive_rules": [
                    {
                        "group_id": f"sg-{i:08d}",
                        "rule": "0.0.0.0/0:22",
                        "risk_level": "high"
                    }
                    for i in range(2)
                ]
            },
            compliance_summary={
                "total_resources": 30,
                "compliant_resources": 10,  # Every 3rd resource
                "non_compliant_resources": 20,
                "compliance_percentage": 33.3
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 30,
                "processing_time_seconds": 4.2
            }
        )
        
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            filename_template="large_dataset_test_{timestamp}"
        )
        
        generator = DocumentGenerator(config)
        summary = generator.generate_bom_documents(large_bom_data)
        
        # Should succeed
        self.assertEqual(summary.successful_formats, 1)
        result = summary.results[0]
        self.assertTrue(result.success)
        
        # Verify file size is reasonable for large dataset
        output_file = os.path.join(self.temp_dir, result.filename)
        self.assertTrue(os.path.exists(output_file))
        self.assertGreater(os.path.getsize(output_file), 15000)  # Should be larger file
        
        # Verify document structure
        doc = docx.Document(output_file)
        
        # Should have substantial content
        self.assertGreater(len(doc.paragraphs), 30)
        self.assertGreater(len(doc.tables), 5)
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_error_recovery(self):
        """Test Word generation error recovery."""
        # Create invalid BOM data to trigger validation error
        invalid_bom_data = BOMData(
            resources=[],  # Empty resources should trigger validation error
            generation_metadata={}
        )
        
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            filename_template="error_test_{timestamp}",
            validate_before_generation=True
        )
        
        generator = DocumentGenerator(config)
        
        # Should raise DocumentGenerationError due to validation failure
        with self.assertRaises(Exception):  # DocumentGenerationError
            generator.generate_bom_documents(invalid_bom_data)
            
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_with_minimal_data(self):
        """Test Word generation with minimal data."""
        minimal_bom_data = BOMData(
            resources=[
                {
                    "service": "EC2",
                    "type": "Instance",
                    "id": "i-minimal",
                    "compliance_status": "compliant"
                }
            ],
            network_analysis={},
            security_analysis={},
            compliance_summary={
                "total_resources": 1,
                "compliant_resources": 1,
                "non_compliant_resources": 0,
                "compliance_percentage": 100.0
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 1
            }
        )
        
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            filename_template="minimal_test_{timestamp}"
        )
        
        generator = DocumentGenerator(config)
        summary = generator.generate_bom_documents(minimal_bom_data)
        
        # Should succeed even with minimal data
        self.assertEqual(summary.successful_formats, 1)
        result = summary.results[0]
        self.assertTrue(result.success)
        
        # Verify file was created
        output_file = os.path.join(self.temp_dir, result.filename)
        self.assertTrue(os.path.exists(output_file))
        
        # Verify basic structure
        doc = docx.Document(output_file)
        self.assertGreater(len(doc.paragraphs), 5)
        
        # Should have at least basic sections
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn("Executive Summary", " ".join(headings))
        
    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx not available")
    def test_word_service_sections(self):
        """Test Word generation with multiple services."""
        multi_service_resources = []
        services = ["EC2", "S3", "RDS", "LAMBDA"]
        
        for service in services:
            for i in range(2):  # 2 resources per service
                multi_service_resources.append({
                    "service": service,
                    "type": f"Resource{i}",
                    "id": f"{service.lower()}-{i:03d}",
                    "name": f"{service} Resource {i}",
                    "region": "us-east-1",
                    "compliance_status": "compliant" if i % 2 == 0 else "non_compliant"
                })
                
        multi_service_bom_data = BOMData(
            resources=multi_service_resources,
            network_analysis={"total_vpcs": 2, "total_subnets": 4},
            security_analysis={"total_security_groups": 5, "high_risk_rules": 1},
            compliance_summary={
                "total_resources": 8,
                "compliant_resources": 4,
                "non_compliant_resources": 4,
                "compliance_percentage": 50.0
            },
            generation_metadata={
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "total_resources": 8
            }
        )
        
        config = DocumentConfig(
            output_formats=["word"],
            output_directory=self.temp_dir,
            filename_template="multi_service_test_{timestamp}"
        )
        
        generator = DocumentGenerator(config)
        summary = generator.generate_bom_documents(multi_service_bom_data)
        
        self.assertTrue(summary.results[0].success)
        
        # Load and check for service sections
        output_file = os.path.join(self.temp_dir, summary.results[0].filename)
        doc = docx.Document(output_file)
        
        # Should have sections for each service
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # Check for service-specific headings
        for service in services:
            service_found = any(service in heading for heading in headings)
            self.assertTrue(service_found, f"{service} section not found in document")


if __name__ == '__main__':
    unittest.main()